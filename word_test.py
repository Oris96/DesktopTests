from clicknium import clicknium as cc, locator, ui
from clicknium.common.enums import * 
from AppOpener import open
import pytest
import docx
import os


file_name_default = "Document1"
file_names_list = [file_name_default]
file_path = os.getcwd()


def save_changes():
    ui(locator.winword.control_panel.file, {"filename":file_name_default}).click()
    ui(locator.winword.listitem_save_as).click()
    ui(locator.winword.button_browse).click()
    ui(locator.winword.window_save_as.edit_file_name).clear_text("send-hotkey")
    ui(locator.winword.window_save_as.edit_file_name).set_text(file_path + os.sep + file_names_list[-1])
    ui(locator.winword.window_save_as.button_save).click()

# Check if document with current name already exist rewrite it 
    if cc.is_existing(locator.winword.window_save_before_close_word.rewrite_file_options_window):
        ui(locator.winword.window_save_before_close_word.rewrite_file_options_window).send_hotkey("{ENTER}")


@pytest.fixture()
def open_close_ms_word():
    
    open("word")

# Check if word opened without blank document then create it
    if cc.is_existing(locator.winword.button_create_new_document):
        ui(locator.winword.button_create_new_document).click()
        for role in ["pane", "group"]:
            variables = {"role":role}
            if cc.is_existing(locator.winword.new_blank_document, variables):
                ui(locator.winword.new_blank_document, variables).double_click()
                break

    yield

# Try to close word. If file were saved with new file name uses current file name   
    for name in file_names_list:
        variables = {"filename":name}
        if cc.is_existing(locator.winword.control_panel.file, variables):
            ui(locator.winword.control_panel.file, variables).send_hotkey("%{F4}")
            break

# If word document wasn't saved before, close Word app without saving
    if cc.wait_appear(locator.winword.window_save_before_close_word.button_dont_save, wait_timeout=2):
        ui(locator.winword.window_save_before_close_word.button_dont_save).click()


def test_create_and_save_file(open_close_ms_word):    
    file_name = "CreateAndSave.docx"
    file_names_list.append(file_name + ".odt")
    file_names_list.append(file_name)
    text = "Hello, "

    ui(locator.winword.body.edit_body).set_text(text)

    save_changes()

    actual_result = docx.Document(file_name).paragraphs[0].text
    assert actual_result == text


def test_create_table_and_edit_style(open_close_ms_word):
    file_name = "CreateTable.docx"
    file_names_list.append(file_name + ".odt")
    file_names_list.append(file_name)
    text = "Create a table"

    ui(locator.winword.body.edit_body).set_text("Create a table")

# Create table and merge 1st row
    ui(locator.winword.body.edit_body).click()
    ui(locator.winword.control_panel.insert).click()
    ui(locator.winword.control_panel.Insert.table).drag_drop(50, 180)
    ui(locator.winword.body.col1_row1_table_cell).send_hotkey("+({RIGHT 4})", "click")
    ui(locator.winword.control_panel.layout).click()
    ui(locator.winword.control_panel.Layout.merge_cells).click()

# Edit style
    ui(locator.winword.control_panel.home).click()
    ui(locator.winword.control_panel.Home.select).click()
    ui(locator.winword.control_panel.Home.select_all).click()
    ui(locator.winword.control_panel.Home.button_bold).click()
    ui(locator.winword.control_panel.Home.button_center).click()
    ui(locator.winword.control_panel.Home.button_italic).click()
    ui(locator.winword.control_panel.Home.edit_font_size).send_hotkey("20{ENTER}")

    save_changes()

    actual_result = docx.Document(file_name).paragraphs[0].text
    assert actual_result == text
